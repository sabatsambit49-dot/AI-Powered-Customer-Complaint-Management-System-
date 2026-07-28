import sys
import os

def create_docx(filename, content_dict):
    try:
        import docx
        doc = docx.Document()
        doc.add_heading(content_dict['title'], level=1)
        for section, text in content_dict['sections'].items():
            doc.add_heading(section, level=2)
            doc.add_paragraph(text)
        doc.save(filename)
        print(f"Created {filename}")
    except ImportError:
        # Fallback to simple text file if docx library isn't installed yet
        txt_name = filename.replace('.docx', '.txt')
        with open(txt_name, 'w', encoding='utf-8') as f:
            f.write(content_dict['title'] + "\n\n")
            for section, text in content_dict['sections'].items():
                f.write(f"=== {section} ===\n{text}\n\n")
        print(f"Fallback: Created {txt_name}")

def create_pdf(filename, content_dict):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        
        doc = SimpleDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        story.append(Paragraph(content_dict['title'], styles['Heading1']))
        story.append(Spacer(1, 12))
        
        for section, text in content_dict['sections'].items():
            story.append(Paragraph(section, styles['Heading2']))
            story.append(Spacer(1, 6))
            story.append(Paragraph(text, styles['Normal']))
            story.append(Spacer(1, 12))
            
        doc.build(story)
        print(f"Created {filename}")
    except ImportError:
        # Fallback to simple text file if reportlab isn't installed yet
        txt_name = filename.replace('.pdf', '.txt')
        with open(txt_name, 'w', encoding='utf-8') as f:
            f.write(content_dict['title'] + "\n\n")
            for section, text in content_dict['sections'].items():
                f.write(f"=== {section} ===\n{text}\n\n")
        print(f"Fallback: Created {txt_name}")

if __name__ == '__main__':
    doc1 = {
        'title': 'CUSTOMER QUALITY COMPLAINT REPORT - DISCOLORED CAPSULES',
        'sections': {
            '1. Origin & Customer Details': 'Complaint Source: Wholesale Distributor (Apex Health Supplies)\nCustomer Name: Apex Health Wholesalers (Contact: David Miller)\nCustomer Email: d.miller@apexhealth.com\nComplaint Date: 2026-07-21',
            '2. Product & Batch Identification': 'Product Name: Omeprazole Delayed-Release Capsules\nProduct Strength: 20mg USP\nBatch/Lot Number: OMP-5510\nManufacturing Date: 2025-11-10\nExpiry Date: 2027-11-09\nQuantity Affected: 12 bottles (1,200 capsules total)',
            '3. Complaint Details': 'Complaint Type: Physical Appearance Defect / Discoloration\nDetailed Description: Customer returned 12 unopened sealed bottles of Omeprazole 20mg Capsules Batch OMP-5510. Upon opening sample bottle for QA intake inspection, capsules displayed mottled dark brown spot discoloration instead of standard opaque white/blue capsule shells. Internal humidity indicator card inside secondary bottle had turned pink, suggesting moisture ingress during packaging or defective bottle cap induction sealing.'
        }
    }
    
    doc2 = {
        'title': 'QUALITY CONTROL INVESTIGATION: OUT OF SPECIFICATION POTENCY ASSAY',
        'sections': {
            '1. Origin & Customer Details': 'Complaint Source: Contract Testing Laboratory (PharmaAnalytica Inc)\nCustomer Name: PharmaAnalytica Quality Audit Division (Auditor: Dr. Karen Vance)\nCustomer Email: kvance@pharmaanalytica.org\nComplaint Date: 2026-07-24',
            '2. Product & Batch Identification': 'Product Name: Atorvastatin Calcium Tablets\nProduct Strength: 40mg\nBatch/Lot Number: ATV-8833\nManufacturing Date: 2026-02-28\nExpiry Date: 2028-02-27\nQuantity Affected: 1 commercial lot (50,000 tablets)',
            '3. Complaint Details': 'Complaint Type: Out of Specification (OOS) Potency / Sub-potent Batch\nDetailed Description: Third-party verification testing of stability sample Batch ATV-8833 returned Atorvastatin active assay value of 86.4% of label claim (USP Acceptance Criteria: 90.0% - 110.0%). HPLC chromatogram indicated significant degradation peak corresponding to Atorvastatin lactone impurity. Batch failed 6-month accelerated stability testing. Requesting immediate recall protocol evaluation and root cause analysis of raw material degradation.'
        }
    }

    os.makedirs('sample-data', exist_ok=True)
    create_docx('sample-data/complaint_06_discolored_capsules.docx', doc1)
    create_pdf('sample-data/complaint_07_out_of_spec_potency.pdf', doc2)
